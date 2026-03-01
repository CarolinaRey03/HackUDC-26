import io
import os
import uuid
import pandas as pd
from typing import Optional
from fastapi import HTTPException, UploadFile, status

from langdetect import detect

from PyPDF2 import PdfReader
from sympy import content
from sentence_transformers import SentenceTransformer
from app.config import settings
from app.core.logging import setup_logger
from app.repositories.documents_repo import (
    analyze_text_es,
    delete_document_es,
    get_document_metadata_es,
    index_document_es,
    list_documents_es,
    search_documents_es,
)
from app.schemas.documents_schemas import DocumentInfo
from odf.opendocument import load
from odf import meta, teletype, dc
from odf.text import P
from docx import Document

_logger = setup_logger(__name__)
_model: SentenceTransformer | None = None


def get_model() -> SentenceTransformer:
    global _model
    if _model is None:
        _logger.debug("Loading embedding model: %s", settings.embedding_model)
        _model = SentenceTransformer(settings.embedding_model)
    return _model


async def index_document(file: UploadFile) -> None:
    content = await file.read()
    filename = file.filename or ""
    content_type = file.content_type or ""

    text = _extract_text(content, content_type, filename)
    language = detect(text)
    chunks = _chunk_text(text)
    
    # Metadata extraction
    title = None
    author = None
    category = "Other"
    
    ct = content_type.lower()
    name_lower = filename.lower()
    
    if ct == "application/pdf" or name_lower.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(content))
            # Rename to avoid shadowing the 'meta' import
            pdf_info = reader.metadata
            title = pdf_info.get("/Title")
            author = pdf_info.get("/Author")
            
            # Additional fallback for common keys
            if not title: title = pdf_info.get("/Subject")
            if not author: author = pdf_info.get("/Creator")
            
            category = "Document"
        except Exception as e:
            _logger.error(f"Error extracting PDF metadata: {e}")
            pass
    elif "wordprocessingml" in ct or name_lower.endswith(".docx"):
        try:
            doc = Document(io.BytesIO(content))
            prop = doc.core_properties
            title = prop.title if prop.title else None
            author = prop.author if prop.author else None
            category = "Document"
        except Exception:
            pass
    elif "spreadsheet" in ct or name_lower.endswith((".xls", ".xlsx", ".csv", ".ods")):
        category = "Spreadsheet"
    elif "text/plain" in ct or name_lower.endswith(".txt"):
        category = "Text"
    elif "oasis.opendocument.text" in ct or name_lower.endswith((".odt", ".odf")):
        try:
            odf_doc = load(io.BytesIO(content))
            titles = odf_doc.meta.getElementsByType(dc.Title)
            creators = odf_doc.meta.getElementsByType(dc.Creator)
            initial_creators = odf_doc.meta.getElementsByType(meta.InitialCreator)
            
            title = teletype.get_text(titles[0]) if titles else None
            author = teletype.get_text(creators[0]) if creators else \
                     (teletype.get_text(initial_creators[0]) if initial_creators else None)
            category = "Document"
        except Exception:
            pass

    _logger.debug(
        "Chunks: %d — preview:\n%s", len(chunks), "\n".join(text.splitlines()[:6])
    )

    analyzed_chunks = [analyze_text_es(chunk, language) for chunk in chunks]
    embeddings = get_model().encode(analyzed_chunks).tolist()

    file_id = str(uuid.uuid4())
    os.makedirs(settings.files_dir, exist_ok=True)
    file_path = os.path.join(settings.files_dir, file_id)
    with open(file_path, "wb") as f:
        f.write(content)

    index_document_es(
        file_id=file_id,
        filename=filename,
        content_type=content_type,
        language=language,
        chunks=analyzed_chunks,
        embeddings=embeddings,
        title=title,
        author=author,
        category=category
    )


def _chunk_text(text: str, max_chars: int = 500) -> list[str]:
    paragraphs = [p.strip() for p in text.splitlines() if p.strip()]
    chunks: list[str] = []
    current = ""
    for p in paragraphs:
        if len(current) + len(p) + 1 <= max_chars:
            current = (current + "\n" + p).strip()
        else:
            if current:
                chunks.append(current)
            current = p
    if current:
        chunks.append(current)
    return chunks


def list_documents(limit: int) -> list[DocumentInfo]:
    return [DocumentInfo(**doc) for doc in list_documents_es(limit)]


def search_documents(
    limit: int,
    query: str,
    language: Optional[str] = None,
    type: Optional[str] = None,
    date: Optional[int] = None,
) -> list[DocumentInfo]:
    embedding: list[float] = []

    if query is not None:
        embedding = get_model().encode([query])[0].tolist()
    return [
        DocumentInfo(**doc)
        for doc in search_documents_es(
            embedding, limit, language=language, file_type=type, date=date
        )
    ]


def delete_document(file_id: str) -> None:
    if not delete_document_es(file_id):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Document not found"
        )

    file_path = os.path.join(settings.files_dir, file_id)
    if os.path.exists(file_path):
        os.remove(file_path)


def get_document(file_id: str) -> tuple[str, str, str]:
    """Returns (file_path, filename, content_type)."""
    meta = get_document_metadata_es(file_id)
    if meta is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Document not found"
        )

    file_path = os.path.join(settings.files_dir, file_id)
    if not os.path.exists(file_path):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="File not found on disk"
        )

    return file_path, meta["filename"], meta["content_type"]

def get_all_documents_filtered(limit: int, language: Optional[str] = None,
    type: Optional[str] = None,
    date: Optional[int] = None) -> list[dict]:
    return list_documents_es(limit, language=language, file_type=type, date=date)

def _extract_text(content: bytes, content_type: str, filename: str) -> str:
    ct = (content_type or "").lower()
    name = (filename or "").lower()

    if ct == "application/pdf" and name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(content))
        text = "\n".join(page.extract_text() for page in reader.pages)
        _logger.debug("\n".join(text.splitlines()[:6]))
        return text

    if (
        ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        and name.endswith(".docx")
    ):
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)

    if "application/vnd.oasis.opendocument.text" in ct or (name.endswith((".odt", ".odf")) and not name.endswith(".ods")):
        odf_doc = load(io.BytesIO(content))
        return "\n".join(str(p) for p in odf_doc.body.getElementsByType(P))

    if ct in ["text/csv", "application/csv"] or name.endswith(".csv"):
        try:
            csv_text = content.decode("utf-8")
        except UnicodeDecodeError:
            csv_text = content.decode("latin1", errors="replace")
            clean_text = csv_text.replace(";", " ").replace(",", " ")

            return clean_text

    if ct == "text/plain" and name.endswith(".txt"):
        return content.decode("utf-8", errors="replace")

    spreadsheet_extensions = (".xlsx", ".xls", ".ods")
    spreadsheet_types = [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", # xlsx
        "application/vnd.ms-excel",                                          # xls
        "application/vnd.oasis.opendocument.spreadsheet"                     # ods
    ]

    if ct in spreadsheet_types or name.endswith(spreadsheet_extensions):

        try:
            engine = 'odf' if name.endswith('.ods') or 'opendocument.spreadsheet' in ct else None

            dfs = pd.read_excel(io.BytesIO(content), sheet_name=None, engine=engine)
            text_parts = []

            for sheet_name, df in dfs.items():
                csv_string = df.to_csv(sep=" ", index=False, header=False, na_rep="")

                lineas_limpias = [linea.strip() for linea in csv_string.splitlines() if linea.strip()]
                if lineas_limpias:
                    text_parts.append("\n".join(lineas_limpias))

            return "\n".join(text_parts)

        except Exception as e:
            _logger.error(f"Error procesando hoja de cálculo con pandas: {e}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST, detail="Error al procesar el archivo Excel/ODS"
            )

    raise HTTPException(
        status_code=status.HTTP_400_BAD_REQUEST, detail="Document format is not valid"
    )